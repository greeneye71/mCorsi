"""Crea il modello DOCX dimostrativo incluso in mCorsi.

Preset: standard_business_brief.
Override nominato ``a4_landscape_certificate``: A4 orizzontale, margini 18 mm,
area 261 x 174 mm, Aptos/Calibri, composizione centrata da credenziale,
palette navy #123F3A e oro #C79A43. L'override è applicato in modo uniforme.
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "mcorsi" / "assets" / "default_certificate.docx"
NAVY = RGBColor(0x12, 0x3F, 0x3A)
GOLD = RGBColor(0xC7, 0x9A, 0x43)
INK = RGBColor(0x20, 0x2B, 0x2A)
MUTED = RGBColor(0x61, 0x6E, 0x6B)


def set_font(run, size, *, color=INK, bold=False, italic=False, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def paragraph(doc, text="", *, size=11, color=INK, bold=False, italic=False, before=0, after=6, align=WD_ALIGN_PARAGRAPH.CENTER):
    item = doc.add_paragraph()
    item.alignment = align
    item.paragraph_format.space_before = Pt(before)
    item.paragraph_format.space_after = Pt(after)
    item.paragraph_format.line_spacing = 1.10
    set_font(item.add_run(text), size, color=color, bold=bold, italic=italic)
    return item


def set_cell_width(cell, width_mm):
    cell.width = Mm(width_mm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    tc_w.set(qn("w:w"), str(round(width_mm / 25.4 * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_page_border(section):
    sect_pr = section._sectPr
    borders = OxmlElement("w:pgBorders")
    borders.set(qn("w:offsetFrom"), "page")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "double")
        border.set(qn("w:sz"), "14")
        border.set(qn("w:space"), "18")
        border.set(qn("w:color"), "123F3A")
        borders.append(border)
    sect_pr.append(borders)


def build():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)
    add_page_border(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_p.paragraph_format.space_after = Pt(0)
    set_font(header_p.add_run("mCorsi  •  FORMAZIONE E COMPETENZE"), 8.5, color=MUTED, bold=True)

    paragraph(doc, "ATTESTATO DI PARTECIPAZIONE", size=12, color=GOLD, bold=True, after=8)
    paragraph(doc, "Si attesta che", size=11, color=MUTED, italic=True, after=4)
    paragraph(doc, "{{ participant_full_name }}", size=28, color=NAVY, bold=True, after=7)
    paragraph(doc, "nato/a a {{ birth_place }} il {{ birth_date }}  •  C.F. {{ tax_code }}", size=10, color=MUTED, after=12)
    paragraph(doc, "ha partecipato e superato la valutazione prevista per il corso", size=11, after=5)
    paragraph(doc, "{{ course_title }}", size=20, color=INK, bold=True, after=7)
    paragraph(doc, "svolto il {{ course_date }}  •  Codice corso {{ course_code }}", size=10.5, color=MUTED, after=6)
    paragraph(doc, "Azienda: {{ company_name }}  •  P. IVA {{ company_vat }}", size=9.5, color=MUTED, after=9)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [80, 101, 80]
    for cell, width in zip(table.rows[0].cells, widths):
        set_cell_width(cell, width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    left, middle, right = table.rows[0].cells
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(p.add_run("Rilascio\n"), 8, color=MUTED, bold=True)
    set_font(p.add_run("{{ issue_date }}\n"), 10, bold=True)
    set_font(p.add_run("Scadenza: {{ expiry_date }}"), 8.5, color=MUTED)
    p = middle.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("{{ signature_image }}"), 9)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(p.add_run("Il responsabile del corso\n"), 8, color=MUTED, bold=True)
    set_font(p.add_run("{{ signer_name }}\n"), 10, bold=True)
    set_font(p.add_run("{{ signer_title }}"), 8.5, color=MUTED)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(0)
    set_font(footer_p.add_run("Attestato n. {{ certificate_number }}  •  Documento generato da mCorsi"), 7.5, color=MUTED)

    doc.core_properties.title = "Modello attestato mCorsi"
    doc.core_properties.subject = "Modello DOCX compilabile da mCorsi"
    doc.core_properties.author = "mCorsi"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
